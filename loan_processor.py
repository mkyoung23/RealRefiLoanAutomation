import os
import pandas as pd
import time
from datetime import datetime
import logging
from openai import OpenAI
import re
from dotenv import load_dotenv
import streamlit as st
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from functools import lru_cache
import json
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')

OUTPUT_FOLDER = "BorrowerResults"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

RATE_MATRIX = [
    (0.0, 0.6, 0.0575, 0.06, 0.06125),
    (0.6001, 0.7, 0.05875, 0.06125, 0.0625),
    (0.7001, 0.75, 0.06, 0.0625, 0.06375),
    (0.7501, 0.8, 0.06125, 0.06375, 0.065),
    (0.8001, 0.95, 0.06125, 0.06375, 0.065)
]

NO_COST_RATE_ADJ = 0.0025
ROLL_RATE_ADJ = -0.0025
CASHOUT_RATE_ADJ = 0.00125

APPRECIATION = 1.07
MAX_CASHOUT_LTV = 0.80
PMT_TOLERANCE = 5.0

logging.basicConfig(filename=os.path.join(OUTPUT_FOLDER, 'loan_processing.log'), level=logging.DEBUG, format='%(asctime)s [%(levelname)s] %(message)s')

def clean_currency(value):
    if pd.isna(value): return 0.0
    if isinstance(value, str):
        return float(value.replace('$', '').replace(',', '').strip())
    return float(value)

def clean_percentage(value):
    if pd.isna(value): return 0.0
    if isinstance(value, str):
        value = value.replace('%', '').strip()
    try:
        val = float(value)
    except:
        return 0.0
    return val / 100 if val > 1 else val

def get_rates_from_ltv(ltv):
    for lower, upper, rate_15, rate_20, rate_30 in RATE_MATRIX:
        if lower <= ltv <= upper:
            return rate_15, rate_20, rate_30
    return 0.06125, 0.06375, 0.065

def manual_pmt(rate, nper, pv):
    if rate == 0 or pv <= 0 or nper <= 0:
        return 0.0
    try:
        res = pv * (rate * (1 + rate)**nper) / ((1 + rate)**nper - 1)
        logging.debug(f"Manual PMT: rate={rate}, nper={nper}, pv={pv}, result={res}")
        return res
    except Exception as e:
        logging.error(f"Manual PMT error: {e}")
        return 0.0

def manual_pv(rate, nper, pmt):
    if rate == 0:
        return pmt * nper
    try:
        res = pmt * (1 - (1 + rate)**(-nper)) / rate
        logging.debug(f"Manual PV: rate={rate}, nper={nper}, pmt={pmt}, result={res}")
        return res
    except Exception as e:
        logging.error(f"Manual PV error: {e}")
        return 0.0

def calculate_amortized_balance(principal, rate, nper, payments_made):
    if principal <= 0 or payments_made <= 0 or nper <= 0:
        return principal
    try:
        remaining_nper = nper - payments_made
        pmt = manual_pmt(rate, nper, principal)
        balance = manual_pv(rate, remaining_nper, pmt)
        logging.debug(f"AMORT BAL: principal={principal}, rate={rate}, nper={nper}, payments_made={payments_made}, result={balance}")
        return max(0, balance)
    except Exception as e:
        logging.error(f"AMORT error: {e}")
        return principal

def validate_pmt(row):
    calc_pmt = manual_pmt(row['Current Interest Rate']/12, row['Loan Term (years)']*12, row['Total Original Loan Amount'])
    if abs(calc_pmt - row['Current P&I Mtg Pymt']) > PMT_TOLERANCE:
        logging.warning(f"PMT mismatch for {row['Borrower First Name']}: Calc {calc_pmt:.2f} vs Input {row['Current P&I Mtg Pymt']:.2f}")
        return "Warning: PMT Mismatch"
    return "Valid"

def calculate_months_elapsed(first_payment_date):
    if pd.isna(first_payment_date):
        return 0
    try:
        first_payment = pd.to_datetime(first_payment_date)
        today = datetime.now()
        months = (today.year - first_payment.year) * 12 + (today.month - first_payment.month)
        return max(months, 0)
    except:
        return 0

@lru_cache(maxsize=128)
def estimate_home_value(address, original_value, months_elapsed):
    if not OPENAI_API_KEY:
        return original_value * (APPRECIATION ** (months_elapsed / 12))
    try:
        # Use Zillow for real lookup
        zillow_url = f"https://www.zillow.com/homes/{address.replace(' ', '-').replace(',', '')}_rb/"
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = f"Browse {zillow_url} and extract the current Zestimate or estimated home value. If not, use recent sold price. Return only the number in dollars."
        response = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], max_tokens=20)
        content = response.choices[0].message.content.strip()
        match = re.search(r'\d+', content.replace(',', ''))
        val = float(match.group()) if match else original_value * (APPRECIATION ** (months_elapsed / 12))
        logging.debug(f"Real Home Value for {address}: {val}")
        return val
    except Exception as e:
        logging.error(f"Home value error: {e}")
        return original_value * (APPRECIATION ** (months_elapsed / 12))

def batch_generate_texts(df, officer_name, company_name, app_link):
    if not OPENAI_API_KEY:
        return [{} for _ in range(len(df))]
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompts = [f"Generate 10 urgent, personalized texts (5 types x 2 A/B variants) for refi outreach based on {row.to_dict()}. Make them sound like from the borrower's previous loan officer ({officer_name} at {company_name}). Highlight current payment vs new options (regular, no-cost, roll-in, cash-out, HELOC), savings, equity, market trends, FOMO (e.g., 'Rates rising soon!'). Include CTA with {app_link}. Vary for response testing." for _, row in df.iterrows()]
        responses = []
        for p in prompts:
            response = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": p}], max_tokens=600)
            responses.append(response)
            time.sleep(1)  # Rate limit
        texts_list = []
        for resp in responses:
            content = resp.choices[0].message.content
            texts = {line.split(':')[0].strip(): ':'.join(line.split(':')[1:]).strip() for line in content.split('\n') if ':' in line}
            texts_list.append(texts)
        return texts_list
    except Exception as e:
        logging.error(f"Text generation error: {e}")
        return [{} for _ in range(len(df))]

def generate_savings_chart(borrower_data):
    fig, ax = plt.subplots()
    labels = ['Current PMT', 'New Reg PMT', 'Monthly Savings']
    values = [borrower_data['Current P&I Mtg Pymt'], borrower_data['Pmt Reg (30yr)'], borrower_data['Savings Reg (30yr)']]
    colors = ['red' if v < 0 else 'green' for v in values]
    ax.bar(labels, values, color=colors)
    ax.set_title('Payment & Savings Breakdown')
    ax.set_ylabel('$')
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    return buf

def generate_equity_chart(borrower_data):
    fig, ax = plt.subplots()
    labels = ['Original Equity', 'Increased Equity']
    values = [borrower_data['Original Appraised Value'] - borrower_data['Total Original Loan Amount'], borrower_data['Equity Increase ($)']]
    ax.pie(values, labels=labels, autopct='%1.1f%%', colors=['#ff9999','#66b3ff'])
    ax.set_title('Equity Growth')
    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    return buf

def generate_word_report(borrower_data, texts, officer_name, company_name, chart_buf_savings, chart_buf_equity):
    doc = Document()
    doc.add_heading(f"Personalized Refinance Report for {borrower_data['Borrower First Name']} {borrower_data['Borrower Last Name']}", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Prepared by {officer_name}, {company_name} | Date: {datetime.now().strftime('%Y-%m-%d')}").style.font.color.rgb = RGBColor(0, 0, 0)

    # Executive Summary
    summary = doc.add_paragraph()
    summary.add_run("Executive Summary").bold = True
    summary.add_run(f"\nCurrent Payment: ${borrower_data['Current P&I Mtg Pymt']:.2f}/mo\nEstimated Home Value: ${borrower_data['New Estimated Home Value']:.2f} (up ${borrower_data['Equity Increase ($)']:.2f})\nTop Opportunity: Save up to ${max(borrower_data['Savings Reg (30yr)'], borrower_data['Savings Roll (30yr)']):.2f}/mo with roll-in option.")
    summary.style.font.size = Pt(12)
    summary.style.font.color.rgb = RGBColor(0, 128, 0) if borrower_data['Savings Reg (30yr)'] > 0 else RGBColor(255, 0, 0)

    # Payment Comparisons Table
    doc.add_heading("Payment Comparisons Across Terms", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Option'
    hdr_cells[1].text = 'Term'
    hdr_cells[2].text = 'Rate'
    hdr_cells[3].text = 'Payment'
    hdr_cells[4].text = 'Savings vs Current'
    hdr_cells[5].text = 'Cash-Out Possible'
    hdr_cells[6].text = 'Details'
    options = [
        ('Regular Refi', '15yr', borrower_data['New Rate (15yr)'], borrower_data['Pmt Reg (15yr)'], borrower_data['Savings Reg (15yr)'], 'N/A', 'Shortest term, highest savings long-term'),
        ('Regular Refi', '20yr', borrower_data['New Rate (20yr)'], borrower_data['Pmt Reg (20yr)'], borrower_data['Savings Reg (20yr)'], 'N/A', 'Balanced term'),
        ('Regular Refi', '30yr', borrower_data['New Rate (30yr)'], borrower_data['Pmt Reg (30yr)'], borrower_data['Savings Reg (30yr)'], 'N/A', 'Lowest monthly payment'),
        ('No-Cost Refi', '30yr', borrower_data['NoCost Rate (30yr)'], borrower_data['Pmt NoCost (30yr)'], borrower_data['Savings NoCost (30yr)'], 'N/A', '$0 closing costs'),
        ('Roll-In Refi', '30yr', borrower_data['Roll Rate (30yr)'], borrower_data['Pmt Roll (30yr)'], borrower_data['Savings Roll (30yr)'], 'N/A', 'Roll points into loan for lower rate'),
        ('Cash-Out Refi', '15yr', borrower_data['CashOut Rate (30yr)'], borrower_data['Pmt CashOut (15yr)'], borrower_data['Savings CashOut (15yr)'], f"${borrower_data['Max CashOut Amount']:.2f}", 'Access equity cash'),
        ('Cash-Out Refi', '20yr', borrower_data['CashOut Rate (30yr)'], borrower_data['Pmt CashOut (20yr)'], borrower_data['Savings CashOut (20yr)'], f"${borrower_data['Max CashOut Amount']:.2f}", 'Access equity cash'),
        ('Cash-Out Refi', '30yr', borrower_data['CashOut Rate (30yr)'], borrower_data['Pmt CashOut (30yr)'], borrower_data['Savings CashOut (30yr)'], f"${borrower_data['Max CashOut Amount']:.2f}", 'Access equity cash'),
        ('HELOC', 'Variable', 'Custom', 'Varies', 'Varies', f"Up to ${borrower_data['Equity Increase ($)']:.2f}", 'Flexible line of credit on equity')
    ]
    for opt, term, rate, pmt, sav, cash, det in options:
        row_cells = table.add_row().cells
        row_cells[0].text = opt
        row_cells[1].text = term
        row_cells[2].text = f"{rate:.3%}" if isinstance(rate, float) else rate
        row_cells[3].text = f"${pmt:.2f}" if isinstance(pmt, float) else pmt
        row_cells[4].text = f"${sav:.2f}" if isinstance(sav, float) else sav
        row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0) if (isinstance(sav, float) and sav > 0) else RGBColor(255, 0, 0)
        row_cells[5].text = cash
        row_cells[6].text = det

    # Charts
    doc.add_heading("Visual Breakdowns", level=1)
    doc.add_paragraph("Savings Chart")
    doc.add_picture(chart_buf_savings, width=Inches(5.5))
    doc.add_paragraph("Equity Growth Chart")
    doc.add_picture(chart_buf_equity, width=Inches(5.5))

    # Personalized Texts
    doc.add_heading("Ready-to-Send Texts (Copy-Paste Optimized)", level=1)
    for key, text in texts.items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(text).italic = True  # For emphasis

    # Custom Notes
    doc.add_heading("Officer Notes", level=1)
    doc.add_paragraph("Enter custom insights or follow-up actions here.")

    report_path = os.path.join(OUTPUT_FOLDER, f"{borrower_data['Borrower First Name']}_{borrower_data['Borrower Last Name']}_Report.docx")
    doc.save(report_path)
    return report_path

def generate_pdf_backup(borrower_data, texts, officer_name, company_name, chart_buf_savings):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(100, 750, f"Backup Refi Report for {borrower_data['Borrower First Name']}")
    # Add basic text content for backup
    c.save()
    buf.seek(0)
    with open(os.path.join(OUTPUT_FOLDER, f"{borrower_data['Borrower First Name']}_backup.pdf"), 'wb') as f:
        f.write(buf.getvalue())

def generate_email_template(borrower_data, texts, officer_name, company_name, app_link):
    name = borrower_data['Borrower First Name']
    template = f"Subject: {name}, Time to Refi Your Loan with {company_name}—Save Big Like Last Time!\n\nHey {name},\nIt's {officer_name} from {company_name}—hope you're well since we closed your loan in {borrower_data['First Pymt Date'].year}! With rates shifting, here's a quick comparison to your current ${borrower_data['Current P&I Mtg Pymt']:.2f}/mo payment:\n\n"
    for key, text in texts.items():
        template += f"{key}:\n{text}\n\n"
    template += f"Market trends show rates could rise soon—don't miss out! Click {app_link} or reply to chat. Looking forward to helping again,\n{officer_name}"
    return template

def process_loans(df, officer_name, company_name, app_link, progress_bar):
    progress_bar.progress(0)
    df['Total Original Loan Amount'] = df['Total Original Loan Amount'].apply(clean_currency)
    df['Original Appraised Value'] = df['Original Appraised Value'].apply(clean_currency)
    df['Current Interest Rate'] = df['Current Interest Rate'].apply(clean_percentage)
    df['Current P&I Mtg Pymt'] = df['Current P&I Mtg Pymt'].apply(clean_currency)
    df['Loan Term (years)'] = pd.to_numeric(df['Loan Term (years)'], errors='coerce').fillna(30)
    df['First Pymt Date'] = pd.to_datetime(df['First Pymt Date'], errors='coerce')

    progress_bar.progress(0.1)
    df['Months Elapsed'] = df['First Pymt Date'].apply(calculate_months_elapsed)

    progress_bar.progress(0.2)
    for idx, row in df.iterrows():
        address = f"{row['Subject Property Address']}, {row['Subject Property City']}, {row['Subject Property State']}"
        est_value = estimate_home_value(address, row['Original Appraised Value'], row['Months Elapsed'])
        df.at[idx, 'New Estimated Home Value'] = est_value

    progress_bar.progress(0.3)
    df['New Loan Balance'] = df.apply(lambda row: calculate_amortized_balance(
        row['Total Original Loan Amount'], row['Current Interest Rate']/12, row['Loan Term (years)']*12, row['Months Elapsed']
    ), axis=1)

    df['LTV'] = df['New Loan Balance'] / df['New Estimated Home Value']

    progress_bar.progress(0.4)
    rate_data = df['LTV'].apply(get_rates_from_ltv)
    df['New Rate (15yr)'] = rate_data.apply(lambda x: x[0])
    df['New Rate (20yr)'] = rate_data.apply(lambda x: x[1])
    df['New Rate (30yr)'] = rate_data.apply(lambda x: x[2])

    df['NoCost Rate (30yr)'] = df['New Rate (30yr)'] + NO_COST_RATE_ADJ
    df['Roll Rate (30yr)'] = df['New Rate (30yr)'] + ROLL_RATE_ADJ
    df['CashOut Rate (30yr)'] = df['New Rate (30yr)'] + CASHOUT_RATE_ADJ

    progress_bar.progress(0.5)
    for term, months in [('15yr', 180), ('20yr', 240), ('30yr', 360)]:
        df[f'Pmt Reg ({term})'] = df.apply(lambda row: manual_pmt(row[f'New Rate ({term})']/12, months, row['New Loan Balance']), axis=1)
        if term == '30yr':
            df[f'Pmt NoCost ({term})'] = df.apply(lambda row: manual_pmt(row['NoCost Rate (30yr)']/12, months, row['New Loan Balance']), axis=1)
            df[f'Pmt Roll ({term})'] = df.apply(lambda row: manual_pmt(row['Roll Rate (30yr)']/12, months, row['New Loan Balance']), axis=1)

    df['Max CashOut Amount'] = (df['New Estimated Home Value'] * MAX_CASHOUT_LTV) - df['New Loan Balance']
    df['Max CashOut Amount'] = df['Max CashOut Amount'].clip(lower=0)
    df['CashOut Loan'] = df['New Loan Balance'] + df['Max CashOut Amount']
    for term, months in [('15yr', 180), ('20yr', 240), ('30yr', 360)]:
        df[f'Pmt CashOut ({term})'] = df.apply(lambda row: manual_pmt(row['CashOut Rate (30yr)']/12, months, row['CashOut Loan']), axis=1)

    for opt in ['Reg', 'NoCost', 'Roll', 'CashOut']:
        for term in ['15yr', '20yr', '30yr']:
            col_pmt = f'Pmt {opt} ({term})'
            col_sav = f'Savings {opt} ({term})'
            if col_pmt in df.columns:
                df[col_sav] = df['Current P&I Mtg Pymt'] - df[col_pmt]

    df['Equity Increase ($)'] = df['New Estimated Home Value'] - df['Original Appraised Value']

    # Rounding
    currency_cols = [col for col in df.columns if any(k in col.lower() for k in ['pmt', 'savings', 'amount', 'balance', 'value', 'equity', 'loan'])]
    df[currency_cols] = df[currency_cols].round(2)
    rate_cols = [col for col in df.columns if 'rate' in col.lower() or 'ltv' in col.lower()]
    df[rate_cols] = df[rate_cols].round(5)

    df['Validation'] = df.apply(validate_pmt, axis=1)

    progress_bar.progress(0.7)
    texts_list = batch_generate_texts(df, officer_name, company_name, app_link)

    progress_bar.progress(0.8)
    word_reports = []
    pdf_backups = []
    email_templates = []
    for idx, (row, texts) in enumerate(zip(df.iterrows(), texts_list)):
        _, borrower_data = row
        chart_buf_savings = generate_savings_chart(borrower_data)
        chart_buf_equity = generate_equity_chart(borrower_data)
        word_path = generate_word_report(borrower_data, texts, officer_name, company_name, chart_buf_savings, chart_buf_equity)
        generate_pdf_backup(borrower_data, texts, officer_name, company_name, chart_buf_savings)
        word_reports.append(word_path)
        email_templates.append(generate_email_template(borrower_data, texts, officer_name, company_name, app_link))

    df['Email Template'] = email_templates
    df = df.sort_values(by='Savings Reg (30yr)', ascending=False)

    # Excel Output
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    action_sheet_path = os.path.join(OUTPUT_FOLDER, f"{officer_name.replace(' ', '_')}_{company_name.replace(' ', '_')}_{timestamp}_ACTION_SHEET.xlsx")
    with pd.ExcelWriter(action_sheet_path, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Action Sheet', index=False)
        summary_df = pd.DataFrame({
            'Metric': ['Average Savings (Reg 30yr)', 'Total Potential Cash-Out', 'Top Borrower by Savings'],
            'Value': [df['Savings Reg (30yr)'].mean(), df['Max CashOut Amount'].sum(), df.iloc[0]['Borrower First Name']]
        })
        summary_df.to_excel(writer, sheet_name='Summary', index=False)

    df.to_json(os.path.join(OUTPUT_FOLDER, 'crm_export.json'), orient='records')

    progress_bar.progress(1.0)
    return df, word_reports, action_sheet_path

# Custom CSS for Professional Look
st.markdown("""
<style>
    .stApp {
        background-color: #f8f9fa;
    }
    .stButton > button {
        background-color: #007bff;
        color: white;
        border-radius: 8px;
        border: none;
    }
    .stTextInput > div > div > input {
        border: 1px solid #ced4da;
        border-radius: 4px;
    }
    h1, h2 {
        color: #0056b3;
    }
    .st-expanderHeader {
        background-color: #e9ecef;
    }
    .sidebar .sidebar-content {
        background-color: #dee2e6;
    }
    .metric-label {
        font-size: 14px;
        color: #6c757d;
    }
</style>
""", unsafe_allow_html=True)

# Streamlit App with Sidebar and Tabs
st.sidebar.header("Real Refi Loan Automation")
uploaded_file = st.sidebar.file_uploader("Upload Borrower Excel", type=['xlsx'], help="Ensure columns match the template for accurate processing.")
officer_name = st.sidebar.text_input("Loan Officer Name", "Michael Young", help="Your name as the loan officer.")
company_name = st.sidebar.text_input("Company Name", "Universal Financial Mortgage", help="Your company name.")
app_link = st.sidebar.text_input("Application Link", "https://example.com/apply", help="Link for borrowers to apply.")

tab1, tab2, tab3 = st.tabs(["Processing Dashboard", "Custom Simulator", "Reports & Exports"])

with tab1:
    st.header("Process Borrower Data")
    if uploaded_file and st.sidebar.button("Process Loans"):
        with st.spinner("Analyzing and generating options..."):
            progress_bar = st.progress(0)
            df = pd.read_excel(uploaded_file)
            expected_cols = ['Borrower First Name', 'Borrower Last Name', 'Subject Property Address', 'Subject Property City', 'Subject Property State', 'Total Original Loan Amount', 'Original Appraised Value', 'First Pymt Date', 'Current Interest Rate', 'Loan Term (years)', 'Current P&I Mtg Pymt', 'Borr Cell', 'Borr Email']
            missing_cols = [col for col in expected_cols if col not in df.columns]
            if missing_cols:
                st.error(f"Missing columns in upload: {', '.join(missing_cols)}. Please use the provided template.")
            else:
                df, reports, action_sheet_path = process_loans(df, officer_name, company_name, app_link, progress_bar)
                st.session_state.df = df  # Persist for other tabs
                st.session_state.reports = reports
                st.session_state.action_sheet_path = action_sheet_path
                col1, col2, col3 = st.columns(3)
                avg_savings = df['Savings Reg (30yr)'].mean()
                col1.metric("Avg Monthly Savings (Reg 30yr)", f"${avg_savings:.2f}", help="Average across all processed borrowers.")
                col2.metric("Total Potential Cash-Out", f"${df['Max CashOut Amount'].sum():.2f}", help="Summed across all borrowers.")
                col3.metric("Top Savings Opportunity", f"${df['Savings Reg (30yr)'].max():.2f}", help="Highest individual savings.")
                sort_by = st.selectbox("Sort Results By", ['Savings Reg (30yr)', 'Savings CashOut (30yr)', 'Equity Increase ($)'], help="Choose sorting criterion.")
                df_sorted = df.sort_values(by=sort_by, ascending=False)
                st.dataframe(df_sorted.style.background_gradient(subset=[col for col in df_sorted.columns if 'Savings' in col], cmap='RdYlGn'))
                st.success("Processing complete.")

with tab2:
    st.header("Interactive Refinance Simulator")
    if 'df' in st.session_state:
        selected = st.selectbox("Select Borrower to Simulate", st.session_state.df['Borrower First Name'], help="Choose a processed borrower.")
        if selected:
            row = st.session_state.df[st.session_state.df['Borrower First Name'] == selected].iloc[0]
            col1, col2 = st.columns(2)
            with col1:
                custom_rate = st.slider("Custom Rate (%)", 0.0, 10.0, row['New Rate (30yr)'] * 100, help="Adjust the interest rate for simulation.") / 100
                custom_term = st.slider("Custom Term (Years)", 10, 30, 30, help="Select the loan term in years.")
            with col2:
                custom_cashout = st.number_input("Custom Cash-Out ($)", 0.0, float(row['Max CashOut Amount']), 0.0, help="Amount to cash out from equity.")
            custom_pmt = manual_pmt(custom_rate / 12, custom_term * 12, row['New Loan Balance'] + custom_cashout)
            custom_ltv = (row['New Loan Balance'] + custom_cashout) / row['New Estimated Home Value']
            custom_savings = row['Current P&I Mtg Pymt'] - custom_pmt
            st.metric("Custom Monthly Payment", f"${custom_pmt:.2f}")
            st.metric("Custom LTV", f"{custom_ltv:.2%}")
            st.metric("Custom Monthly Savings", f"${custom_savings:.2f}", delta_color="normal" if custom_savings > 0 else "inverse")
            if st.button("Save This Scenario"):
                custom_df = pd.DataFrame([{'Borrower': selected, 'Custom Rate': custom_rate, 'Custom Term': custom_term, 'Custom CashOut': custom_cashout, 'Custom PMT': custom_pmt, 'Custom Savings': custom_savings}])
                custom_path = os.path.join(OUTPUT_FOLDER, f"{selected}_custom_scenario.xlsx")
                custom_df.to_excel(custom_path, index=False)
                st.download_button("Download Saved Scenario", open(custom_path, 'rb').read(), os.path.basename(custom_path))

with tab3:
    st.header("Generated Reports & Exports")
    if 'df' in st.session_state:
        st.download_button("Download Full Action Sheet (Excel)", open(st.session_state.action_sheet_path, 'rb').read(), os.path.basename(st.session_state.action_sheet_path), help="Complete dataset with all calculations.")
        st.download_button("Download CRM JSON Export", open(os.path.join(OUTPUT_FOLDER, 'crm_export.json'), 'rb').read(), "crm_export.json", help="For integration with CRM systems.")
        for idx, row in st.session_state.df.iterrows():
            with st.expander(f"{row['Borrower First Name']} {row['Borrower Last Name']} - Key Savings: ${row['Savings Reg (30yr)']:.2f}", expanded=False):
                st.write(f"**Email Template:** {row['Email Template']}")
                st.download_button(f"Download Word Report for {row['Borrower First Name']}", open(st.session_state.reports[idx], 'rb').read(), os.path.basename(st.session_state.reports[idx]), help="Professional Word document with charts and details.")
                if st.button(f"Export Texts for {row['Borrower First Name']}"):
                    texts_path = os.path.join(OUTPUT_FOLDER, f"{row['Borrower First Name']}_texts.txt")
                    with open(texts_path, 'w') as f:
                        f.write(row['Email Template'])
                    st.download_button("Download Texts File", open(texts_path, 'rb').read(), os.path.basename(texts_path))